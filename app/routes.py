from flask import render_template, request, redirect, url_for, jsonify, flash, session
from app import app, db
from app.models import *
from pydub import AudioSegment
import io
import speech_recognition as sr
from difflib import SequenceMatcher
from docx import Document
from app.decorators import login_required
import json
import google.generativeai as genai
from werkzeug.utils import secure_filename

app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024
app.config['UPLOAD_FOLDER'] = 'app/static/images'
app.config['ALLOWED_EXTENSIONS'] = {'png', 'jpg', 'jpeg', 'gif'}
app.config['TEMPLATES_AUTO_RELOAD'] = True
app.debug = True

# --------------------------------------------------------------Cấu hình API key-------------
genai.configure(api_key="AIzaSyApgUbygeBFmX7cXXIRdRmH9fo7d7FCZec")

@app.route('/score_answer', methods=['POST'])
def score_answer():
    data = request.json
    question = data.get('question', '')
    hints = data.get('hints', [])
    answer = data.get('answer', '')

    if not answer:
        return jsonify({"error": "No answer provided"}), 400

    # Tạo prompt gửi đến Gemini
    # Prompt yêu cầu mô hình trả về JSON chứa score (0-10) và feedback
    prompt = f"""
You are a scoring expert. I will give you a question, hints, and a user answer.
Please consider how well the user's answer matches the question and the hints.
Give a score from 0 to 10 (integer) and a short feedback (1-2 sentences).
Output your answer in JSON format:

{{
  "score": <integer from 0 to 10>,
  "feedback": "<short feedback>"
}}

Question: {question}
Hints: {", ".join(hints)}
User Answer: {answer}
"""

    try:
        # Gọi Gemini API
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(prompt)

        # response.text chứa nội dung trả về của Gemini
        # Thử parse JSON từ text
        # Nếu Gemini trả về đúng JSON thì ok, nếu không cần xử lý ngoại lệ.
        result_text = response.text.strip()

        # Thử loại bỏ ký tự thừa và parse JSON:
        # Nếu mô hình trả về JSON chuẩn, ta có thể parse trực tiếp:
        # Ví dụ output mong đợi:
        # {
        #   "score": 8,
        #   "feedback": "Your answer is good but could be more detailed."
        # }

        # Tìm vị trí { và } để cắt chuỗi (phòng trường hợp model thêm text ngoài).
        start_idx = result_text.find('{')
        end_idx = result_text.rfind('}')
        if start_idx != -1 and end_idx != -1:
            json_str = result_text[start_idx:end_idx+1]
        else:
            # Nếu không tìm thấy cặp {} thì trả về lỗi
            return jsonify({"error": "Could not parse JSON from the model response"}), 500

        # Parse JSON
        parsed = json.loads(json_str)

        score = parsed.get("score")
        feedback = parsed.get("feedback")

        if score is None or feedback is None:
            return jsonify({"error": "JSON missing score or feedback"}), 500

        return jsonify({"score": score, "feedback": feedback}), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500
# -------------------------------------------------------------------



# Trang chủ
@app.route('/')
def index():
    return render_template('index.html')

#-----------------------------Đăng kí---------------------------------------
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']
        confirm_password = request.form['confirm_password']

        # Kiểm tra dữ liệu nhập
        if password != confirm_password:
            flash("Passwords do not match!", "danger")
            return render_template('register.html')

        # Kiểm tra email hoặc username đã tồn tại
        if User.query.filter_by(email=email).first():
            flash("Email already exists!", "danger")
            return render_template('register.html')
        if User.query.filter_by(username=username).first():
            flash("Username already exists!", "danger")
            return render_template('register.html')

        # Tạo user mới
        new_user = User(username=username, email=email)
        new_user.set_password(password)
        db.session.add(new_user)
        db.session.commit()

        flash("Registration successful! Please login.", "success")
        return redirect(url_for('login'))
    return render_template('register.html')

#------------------- Đăng nhập---------------------------------------
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        user = User.query.filter_by(email=email).first()
        
        if user and user.check_password(password):
            session['user_id'] = user.id
            session['username'] = user.username
            session['is_admin'] = user.is_admin  # Lưu thông tin admin vào session

            # Kiểm tra xem người dùng có phải là admin không
            if user.is_admin:
                flash('Login successful! Welcome, Admin.', 'success')
                return redirect(url_for('admin.admin_dashboard'))  # Điều hướng đến trang admin nếu là admin
            else:
                flash('Login successful!', 'success')
                return redirect(url_for('index'))  # Điều hướng đến trang chính nếu không phải admin

        else:
            flash('Invalid email or password.', 'danger')
    return render_template('login.html')

#------------------- Đăng xuất---------------------------------------
@app.route('/logout')
def logout():
    session.pop('user_id', None)
    session.pop('username', None)
    flash('You have been logged out.', 'info')
    return redirect(url_for('login'))

#--------------------------------------- Profile---------------------------------------
@app.route('/profile')
@login_required

def profile():
    if not session.get('user_id'):
        return redirect(url_for('login'))
    user = User.query.get(session.get('user_id'))
    return render_template('profile.html', user=user)
#-----------Edit Profile---------------------------------------
@login_required
@app.route('/profile/edit', methods=['GET', 'POST'])
def profile_edit():
    if 'user_id' not in session:
        flash('You must be logged in to access this page.', 'error')
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])

    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']
        confirm_password = request.form['confirm_password']
        
        # Handle Avatar Update
        if 'image_url' in request.files:
            image = request.files['image_url']
            if image:
                filename = secure_filename(image.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                image.save(filepath)
                user.image_url = filename

        # Handle Password Update
        if password and password != confirm_password:
            flash('Passwords do not match.', 'error')
            return render_template('profile_edit.html', user=user)

        user.username = username
        user.email = email
        if password:
            user.password = generate_password_hash(password)

        db.session.commit()
        flash('Profile updated successfully!', 'success')
        return redirect(url_for('profile_edit'))

    return render_template('profile_edit.html', user=user)


# Trang chủ đề
@app.route('/topic')
@login_required

def topics():
    topics = Topic.query.all()
    return render_template('topic.html', topics=topics)

@app.route('/topic/<int:topic_id>/lessons')
@login_required
def view_lessons(topic_id):
    user_id = session['user_id']  # Lấy user_id từ session
    topic = Topic.query.get_or_404(topic_id)  # Lấy thông tin của Topic

    # Lấy tất cả bài học thuộc topic
    lessons = Lesson.query.filter_by(topic_id=topic_id).all()

    # Lấy danh sách các bài học đã hoàn thành từ bảng Progress
    completed_lessons = Progress.query.filter_by(user_id=user_id, is_learned=1).with_entities(Progress.lesson_id).all()
    completed_lessons_ids = {progress.lesson_id for progress in completed_lessons}

    return render_template(
        'lesson_list.html',
        topic=topic,
        lessons=lessons,
        completed_lessons=completed_lessons_ids
    )

# Bài học
@app.template_filter('to_letter')
def to_letter(number):
    """Chuyển số thành ký tự chữ cái (A, B, C, ...)"""
    return chr(64 + number)

@app.route('/lesson/<int:id>')
@login_required

def lesson_detail(id):
    lesson = Lesson.query.get_or_404(id)
    if lesson.lesson_type == "vocabulary":
        # Nếu là bài học từ vựng, lấy danh sách từ vựng thuộc chủ đề
        vocabularies = Vocabulary.query.filter_by(topic_id=lesson.topic_id).all()
    else:
        vocabularies = None
    questions = Question.query.filter_by(lesson_id=id).all()

    # Lấy tất cả các câu hỏi và các lựa chọn của từng câu hỏi
    question_data = []
    for question in questions:
        choices = Choice.query.filter_by(question_id=question.id).all()
        question_data.append({'question': question, 'choices': choices})

    return render_template('lesson_detail.html', lesson=lesson, vocabularies=vocabularies, question_data=question_data)
    

@app.route('/check_answers', methods=['POST'])
@login_required

def check_answers():
    try:
        data = request.json  # Lấy đáp án từ người dùng
        correct_answers = 0
        total_questions = len(data)
        details = {}

        for question_id, selected_choice_id in data.items():
            question_id = int(question_id.split('-')[1])  # Chuyển từ "question-1" thành ID
            selected_choice_id = int(selected_choice_id)

            # Kiểm tra đáp án đúng
            correct_choice = Choice.query.filter_by(question_id=question_id, is_correct=True).first()
            if correct_choice and correct_choice.id == selected_choice_id:
                correct_answers += 1
                details[question_id] = True
            else:
                details[question_id] = False

        return jsonify({
            "correct": correct_answers == total_questions,
            "score": correct_answers,
            "total": total_questions,
            "details": details,  # Chi tiết từng câu hỏi
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/check_word', methods=['POST'])
def check_word():
    try:
        data = request.json
        user_word = data['word'].strip().lower()  # Chuẩn hóa từ người dùng
        index = data['id']
        db_id = index   # Chuyển từ index (0-based) sang id (1-based)

        # Lấy từ vựng dựa trên db_id
        vocab = Vocabulary.query.get(db_id)
        if not vocab:
            print(f"Vocabulary with id {db_id} (index {index}) not found in the database.")  # Log chi tiết
            return jsonify({"error": "Vocabulary not found"}), 404

        correct_word = vocab.word.strip().lower()  # Chuẩn hóa từ từ CSDL

        # Log chi tiết
        print(f"User word: '{user_word}', ASCII: {[ord(c) for c in user_word]}, Type: {type(user_word)}")
        print(f"Correct word: '{correct_word}', ASCII: {[ord(c) for c in correct_word]}, Type: {type(correct_word)}")
        print(f"Vocabulary id: {db_id}, Correct word from DB: {vocab.word}")

        # So sánh
        if "".join(user_word.split()) == "".join(correct_word.split()):
            return jsonify({"correct": True})
        else:
            return jsonify({
                "correct": False,
                "details": {
                    "user_word": user_word,
                    "correct_word": correct_word,
                    "ascii_user_word": [ord(c) for c in user_word],
                    "ascii_correct_word": [ord(c) for c in correct_word],
                    "type_user_word": str(type(user_word)),
                    "type_correct_word": str(type(correct_word)),
                    "db_index": index,
                    "db_id": db_id
                }
            })

    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": "An error occurred"}), 500

# Chuyển từ mp3 sang text
@app.route('/sound-to-text')
@login_required

def speech_to_text():
    return render_template('sound_to_text.html')

@app.route('/mp3-to-text', methods=['POST'])
@login_required

def mp3_to_text():
    if 'audio' not in request.files:
        return jsonify({"error": "No audio file found"}), 400

    audio_file = request.files['audio']

    try:
        # Đọc file MP3
        audio = AudioSegment.from_mp3(audio_file)
        audio_wav = io.BytesIO()
        audio.export(audio_wav, format="wav")
        audio_wav.seek(0)

        # Sử dụng thư viện SpeechRecognition để chuyển âm thanh thành văn bản
        recognizer = sr.Recognizer()
        with sr.AudioFile(audio_wav) as source:
            audio_data = recognizer.record(source)
            text = recognizer.recognize_google(audio_data)
        
        return jsonify({"text": text})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/add_progress/<int:user_id>/<int:vocabulary_id>', methods=['POST'])
@login_required

def add_progress(user_id, vocabulary_id):
    progress = Progress(user_id=user_id, vocabulary_id=vocabulary_id)
    db.session.add(progress)
    db.session.commit()
    return redirect(url_for('vocabulary'))
    if 'audio' not in request.files:
        return jsonify({"error": "No audio file found"}), 400

    audio_file = request.files['audio']
    file_path = f"static/uploads/{audio_file.filename}"
    audio_file.save(file_path)
    return jsonify({"audio_url": file_path})

# Add Speaking Practice B1 to navigation
@app.context_processor
def inject_nav_links():
    return {
        'nav_links': [
            {"name": "Home", "url": url_for('index')},
            {"name": "Topics", "url": url_for('topics')},
            {"name": "Mp3 to text", "url": url_for('speech_to_text')},
            {"name": "Challenge", "url": url_for('speaking_challenge')}
        ]
    }
# --------------------------------------speaking---------------------------------
# Route hiển thị danh sách topic

@app.route('/speaking', methods=['GET'])
@login_required
def speaking_challenge():
    topics = Topic.query.all()
    return render_template('speaking.html', topics=topics)


@app.route('/speaking/<int:topic_id>', methods=['GET'])
@login_required
def speaking_questions(topic_id):
    # Lấy thông tin topic, câu hỏi, và gợi ý
    topic = Topic.query.get_or_404(topic_id)
    question = Challenge.query.filter_by(topic_id=topic.id).first()
    hints = Hint.query.filter_by(question_id=question.id).all() if question else []

    return render_template(
        'speaking_questions.html',
        topic=topic,
        question=question,
        hints=hints
    )


# Route để lưu câu trả lời vào file docx
import os

@app.route('/save_answer', methods=['POST'])
@login_required
def save_answer():
    data = request.json
    question = data.get('question', '')
    answer = data.get('answer', '')

    # Tạo file .docx
    doc = Document()
    doc.add_heading('Speaking Challenge', level=1)
    doc.add_heading('Question:', level=2)
    doc.add_paragraph(question)
    doc.add_heading('Your Answer:', level=2)
    doc.add_paragraph(answer)

    # Lưu file vào thư mục static/documents
    file_path = os.path.join('static/documents', 'speaking_answer.docx')
    doc.save(file_path)

    return jsonify({
        "message": "Answer saved successfully!",
        "file_path": file_path
    })
# ----------------------------------------------------------Game------------------------------------------
@login_required
@app.route('/game')
def game():
    return render_template('game.html')
# --------------------------------------------------------Progress-----------------------------------------
@app.route('/lesson/<int:lesson_id>', methods=['GET'])
@login_required
def lesson_details(lesson_id):
    user_id = session['user_id']
    lesson = Lesson.query.get_or_404(lesson_id)

    # Lấy dữ liệu từ Progress
    progress = Progress.query.filter_by(user_id=user_id, lesson_id=lesson_id).first()
    is_completed = progress.is_learned if progress else False

    # Lấy từ vựng và câu hỏi cho bài học
    vocabularies = Vocabulary.query.filter_by(lesson_id=lesson_id).all()
    questions = Question.query.filter_by(lesson_id=lesson_id).all()

    # Trả lời đúng (nếu đã làm)
    correct_answers = {}
    if progress and progress.is_learned:
        for question in questions:
            correct_choice = Choice.query.filter_by(question_id=question.id, is_correct=True).first()
            if correct_choice:
                correct_answers[question.id] = correct_choice.id

    return render_template(
        'lession_detail.html',
        lesson=lesson,
        vocabularies=vocabularies,
        question_data=[{'question': q, 'choices': Choice.query.filter_by(question_id=q.id).all()} for q in questions],
        is_completed=is_completed,
        correct_answers=correct_answers
    )

# Route xử lý submit câu trả lời và cập nhật progress
@app.route('/lesson/<int:lesson_id>/submit', methods=['POST'])
@login_required
def submit_lesson(lesson_id):
    try:
        user_id = session['user_id']
        answers = request.json.get('answers')  # Nhận câu trả lời từ client

        # Kiểm tra và lấy các câu hỏi của bài học
        questions = Question.query.filter_by(lesson_id=lesson_id).all()
        correct_answers = {}
        is_all_correct = True

        # So sánh đáp án của người dùng với đáp án đúng
        for question in questions:
            correct_choice = Choice.query.filter_by(question_id=question.id, is_correct=True).first()
            correct_answers[str(question.id)] = correct_choice.id
            if str(question.id) not in answers or int(answers[str(question.id)]) != correct_choice.id:
                is_all_correct = False

        # Cập nhật trạng thái hoàn thành nếu tất cả đáp án đều đúng
        progress = Progress.query.filter_by(user_id=user_id, lesson_id=lesson_id).first()
        if not progress:
            progress = Progress(user_id=user_id, lesson_id=lesson_id, is_learned=0)
            db.session.add(progress)

        if is_all_correct:
            progress.is_learned = 1
            db.session.commit()
            return jsonify({"is_completed": True, "correct_answers": correct_answers})
        else:
            db.session.commit()
            return jsonify({"is_completed": False, "correct_answers": correct_answers})
    except Exception as e:
        print(f"Error: {e}")  # Ghi log lỗi trên server
        return jsonify({"error": "An internal server error occurred"}), 500
